import PyPDF2
from docx import Document
from docx.shared import Pt

# 打开PDF文件
with open("D:\\github\\openai-translator\\tests\\first_5_page.pdf", "rb") as pdf_file:
    # 创建PDF Reader对象
    pdf_reader = PyPDF2.PdfReader(pdf_file)

    # 创建Word文档
    doc = Document()

    # 定义字体和字号
    header_font = 'Arial'
    header_font_size = Pt(14)
    body_font = 'Calibri'
    body_font_size = Pt(12)

    # 遍历PDF页面
    for page_num in range(len(pdf_reader.pages)):
        # 获取当前页
        page = pdf_reader.pages[page_num]

        # 获取页面中的字体信息
        font_info = page['/Resources']['/Font']

        # 将页面内容提取为文本
        text = page.extract_text()

        # 根据文本特征设置字体和字号
        lines = text.split('\n')
        for line in lines:
            # 判断是否为标题行
            if line.isupper():
                run = doc.add_paragraph().add_run(line)
                font = run.font
                font.name = header_font
                font.size = header_font_size
            else:
                run = doc.add_paragraph().add_run(line)
                font = run.font
                font.name = body_font
                font.size = body_font_size

    # 保存Word文档
    doc.save("D:\\github\\openai-translator\\tests\\first_5_page.docx")
